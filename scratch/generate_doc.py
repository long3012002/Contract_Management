import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run()
    
    fldChar1 = create_element('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = create_element('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = create_element('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_callout(doc, text, title="LƯU Ý QUAN TRỌNG", border_color="003366", bg_color="F0F4F8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=140)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.bold = True
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(0, 51, 102)
    
    run_text = p.add_run(text)
    run_text.font.name = "Segoe UI"
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(51, 51, 51)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def build_docx():
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(43, 43, 43)
    
    # Title Block
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_header.add_run("NGÂN HÀNG THƯƠNG MẠI CỔ PHẦN BẢO HIỂM NÔNG NGHIỆP / COOPBANK\nTRUNG TÂM CÔNG NGHỆ THÔNG TIN\n-----------------------------------\n")
    r_sub.font.name = "Segoe UI"
    r_sub.font.size = Pt(10)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(100, 100, 100)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("TỜ TRÌNH, ĐẶC TẢ KỸ THUẬT & YÊU CẦU TRIỂN KHAI\nHỆ THỐNG QUẢN LÝ DỰ ÁN VÀ HỢP ĐỒNG CNTT (RELEASE V1.0)")
    r_title.font.name = "Segoe UI"
    r_title.font.size = Pt(17)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(18)
    r_desc = p_desc.add_run("Tài liệu kỹ thuật tổng thể phục vụ nghiệm thu & đẩy bản (Deploy) lần đầu tại CoopBank")
    r_desc.font.italic = True
    r_desc.font.size = Pt(11)
    r_desc.font.color.rgb = RGBColor(80, 80, 80)

    # Info Box Table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(info_table, "CCCCCC")
    info_data = [
        ("Tên dự án:", "Hệ thống Quản lý Dự án và Hợp đồng CNTT (QLDA)"),
        ("Phiên bản bàn giao:", "Release Version 1.0.0 (Bản chính thức đầu tiên)"),
        ("Đơn vị đề xuất:", "Đội Phát triển Phần mềm - Trung tâm CNTT"),
        ("Đơn vị phối hợp:", "Phòng Quản trị Hạ tầng, An toàn Thông tin, Đội DBA")
    ]
    for idx, (k, v) in enumerate(info_data):
        row = info_table.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        set_cell_background(cell_k, "F5F7FA")
        set_cell_margins(cell_k, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_v, top=80, bottom=80, left=100, right=100)
        
        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        
        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table of Contents Section
    h_toc = doc.add_heading("MỤC LỤC TỰ ĐỘNG", level=1)
    h_toc.style.font.color.rgb = RGBColor(0, 51, 102)
    
    p_toc_note = doc.add_paragraph()
    r_tn = p_toc_note.add_run("(Lưu ý: Mục lục dưới đây tự động cập nhật theo các tiêu đề Heading. Trong Microsoft Word, quý vị có thể nhấn chuột phải vào mục lục và chọn 'Update Field' hoặc bấm phím F9 để cập nhật lại số trang sau khi chỉnh sửa nội dung).")
    r_tn.font.italic = True
    r_tn.font.size = Pt(9.5)
    r_tn.font.color.rgb = RGBColor(120, 120, 120)
    
    add_toc(doc)
    doc.add_page_break()

    # Section 1
    h1 = doc.add_heading("1. Tổng Quan Hệ Thống & Phạm Vi Triển Khai", level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.add_run("Tài liệu này quy định chi tiết các đặc tả kỹ thuật phần mềm, cấu trúc cơ sở dữ liệu, luồng xác thực bảo mật, cùng các danh mục kiểm tra an toàn và quy trình từng bước để triển khai (deploy) bản phát hành đầu tiên cho ứng dụng ")
    r_bold = p.add_run("Quản lý Dự án & Hợp đồng CNTT")
    r_bold.bold = True
    p.add_run(" lên hạ tầng môi trường Staging/Production của CoopBank.")
    
    h1_1 = doc.add_heading("1.1. Phạm vi nghiệp vụ của hệ thống", level=2)
    h1_1.style.font.color.rgb = RGBColor(11, 79, 108)
    
    business_scope = [
        ("Quản lý Dự án (Projects):", "Quản lý danh mục dự án CNTT, tổng ngân sách dự toán, theo dõi trạng thái tiến độ lập kế hoạch, triển khai và hoàn thành."),
        ("Quản lý Gói thầu (Bid Packages):", "Chia nhỏ dự án thành các gói thầu, kiểm soát giá trị dự toán gói thầu và cài đặt ngưỡng cảnh báo chi tiêu (Warning Threshold)."),
        ("Quản lý Hợp đồng (Contracts):", "Lưu trữ hợp đồng ký kết, giá trị hợp đồng, thời hạn hiệu lực, tự động phát tín hiệu nhắc hạn gia hạn hợp đồng (Renewal Reminder)."),
        ("Quản lý Đối tác & Nhà thầu (Partners):", "Lưu trữ thông tin mã số thuế, nhà thầu chính, nhà thầu phụ liên quan trong từng hợp đồng."),
        ("Quản lý Nghị quyết & Văn bản (Resolutions):", "Số hóa văn bản phê duyệt, đính kèm tệp hồ sơ pháp lý liên quan đến dự án và hợp đồng.")
    ]
    for title, desc in business_scope:
        p_c = doc.add_paragraph(style='List Bullet')
        p_c.paragraph_format.space_after = Pt(3)
        rt = p_c.add_run(title + " ")
        rt.bold = True
        p_c.add_run(desc)

    # Section 2: Comprehensive Technical Specification
    h2 = doc.add_heading("2. Đặc Tả Kỹ Thuật Hệ Thống Xây Dựng (System Technical Specification)", level=1)
    h2.style.font.color.rgb = RGBColor(0, 51, 102)
    
    h2_1 = doc.add_heading("2.1. Đặc tả Công nghệ & Kiến trúc Phần mềm Backend", level=2)
    h2_1.style.font.color.rgb = RGBColor(11, 79, 108)
    
    be_tech_table = doc.add_table(rows=7, cols=3)
    be_tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(be_tech_table, "CCCCCC")
    
    be_headers = ["Thành phần (Component)", "Công nghệ / Thư viện", "Mô tả Đặc tả Kỹ thuật"]
    for idx, text in enumerate(be_headers):
        cell = be_tech_table.rows[0].cells[idx]
        set_cell_background(cell, "003366")
        p_h = cell.paragraphs[0]
        r = p_h.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    be_specs = [
        ("Framework Cốt lõi", ".NET 8.0 SDK (C# 12)", "Xây dựng Web API hiệu năng cao trên Kestrel Server, hỗ trợ Dependency Injection native."),
        ("ORM & Database Driver", "EF Core 8.0 + Pomelo MySQL 8.0", "Quản lý truy vấn CSDL qua DbContext, Code-First Migration và LINQ query optimization."),
        ("Xác thực & Bảo mật Token", "Microsoft.AspNetCore.Authentication.JwtBearer", "Xác thực JWT Bearer Token, kiểm tra TokenLifetime, Issuer/Audience và Query Access Token cho WebSocket."),
        ("Realtime Engine", "Microsoft.AspNetCore.SignalR 8.0", "Quản lý kết nối thời gian thực WebSocket, tự động Reconnect và đẩy thông báo qua ReceiveNotification Event."),
        ("Ánh xạ DTO & Mapper", "AutoMapper 12.0", "Tự động ánh xạ dữ liệu giữa Entity CSDL và DTO API (CreateDto, UpdateDto, ResponseDto)."),
        ("Tích hợp Xác thực RADIUS", "RadiusClient Custom Service", "Giao thức UDP 1812 xác thực tài khoản Domain nội bộ Ngân hàng qua Radius Server 10.224.0.94.")
    ]
    for row_idx, data in enumerate(be_specs, start=1):
        row = be_tech_table.rows[row_idx]
        bg = "F9FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p_c = cell.paragraphs[0]
            p_c.paragraph_format.space_after = Pt(0)
            p_c.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    h2_2 = doc.add_heading("2.2. Đặc tả Kiến trúc Mô hình Layered Architecture", level=2)
    h2_2.style.font.color.rgb = RGBColor(11, 79, 108)
    
    p_arch = doc.add_paragraph()
    p_arch.add_run("Mã nguồn Backend được tổ chức phân lớp rõ ràng theo nguyên lý Separation of Concerns (SoC):")
    
    arch_layers = [
        ("Controllers/", "Cửa ngõ tiếp nhận HTTP Request từ Client, thực hiện kiểm tra Filter Audit Log và trả về Response dạng ApiErrorResponse hoặc DTO."),
        ("DTOs/ & Mapper/", "Định nghĩa mẫu dữ liệu đầu vào/đầu ra, tách biệt hoàn toàn giữa Entity nội bộ CSDL và dữ liệu công khai trên API."),
        ("Services/Interfaces/ & Implements/", "Nơi xử lý toàn bộ logic nghiệp vụ (Business Logic), tính toán ngân sách hợp đồng, kiểm tra phân quyền RBAC và thực thi CRUD qua DbCrudService."),
        ("Entity/ & Data/ (AppDbContext)", "Định nghĩa lớp đối tượng kế thừa từ BaseEntity (Id Guid, Code, Name, CreatedAt, UpdatedAt) và cấu hình liên kết bảng Entity Framework Core."),
        ("Middleware/ & Logging/", "Xử lý lỗi tập trung qua ExceptionHandlingMiddleware, bổ sung File Logger và Audit Log Action Filter ghi lại nhật ký tác động."),
        ("Hosted Services (Background Workers):", "Các Tiến trình chạy ngầm: ContractScanWorker (quét hợp đồng hết hạn) và AuditLogRetentionWorker (dọn dẹp nhật ký audit định kỳ).")
    ]
    for layer, desc in arch_layers:
        p_l = doc.add_paragraph(style='List Bullet')
        p_l.paragraph_format.space_after = Pt(3)
        r_l = p_l.add_run(layer + " ")
        r_l.bold = True
        p_l.add_run(desc)

    h2_3 = doc.add_heading("2.3. Đặc tả CSDL & Cơ chế Phân quyền Chi tiết (RBAC)", level=2)
    h2_3.style.font.color.rgb = RGBColor(11, 79, 108)
    
    p_rbac = doc.add_paragraph()
    p_rbac.add_run("Hệ thống sử dụng cơ chế Phân quyền theo Vai trò và Chức năng động (Feature-based Role-Based Access Control):")
    
    p_rbac_1 = doc.add_paragraph(style='List Bullet')
    p_rbac_1.add_run("Bảng Thực thể Phân quyền: ").bold = True
    p_rbac_1.add_run("Bao gồm Users, Roles, Features, UserRoles (nhiều-nhiều) và RolePermissions (RoleId, FeatureId, CanAccess, Permissions string).")
    
    p_rbac_2 = doc.add_paragraph(style='List Bullet')
    p_rbac_2.add_run("Quyền Tối Cao (System Admin): ").bold = True
    p_rbac_2.add_run("Người dùng có 'IsSystemAdmin = true' sẽ tự động vượt qua mọi bước kiểm tra phân quyền.")
    
    p_rbac_3 = doc.add_paragraph(style='List Bullet')
    p_rbac_3.add_run("Cơ chế Phân quyền Chi tiết (Granular Permission): ").bold = True
    p_rbac_3.add_run("Trường 'Permissions' lưu danh sách các thao tác ngăn cách bởi dấu ';' (ví dụ: 'Create;Update;Delete;Approve'). Bộ lọc FeatureAuthorizeFilter tự động kiểm tra tương ứng với phương thức HTTP (GET -> CanAccess, POST -> Create, PUT -> Update, DELETE -> Delete).")

    h2_4 = doc.add_heading("2.4. Đặc tả Cơ chế Bảo mật Multi-Factor Auth (MFA TOTP)", level=2)
    h2_4.style.font.color.rgb = RGBColor(11, 79, 108)
    
    p_mfa = doc.add_paragraph()
    p_mfa.add_run("Hệ thống tích hợp xác thực 2 lớp dựa trên thuật toán TOTP (Time-based One-Time Password - RFC 6238):")
    
    p_mfa_step = doc.add_paragraph()
    p_mfa_step.add_run("1. Khi người dùng bật MFA, hệ thống tạo một chuỗi Bí mật Base32 (TwoFactorSecret) và cung cấp QR Code để quét vào ứng dụng Google/Microsoft Authenticator.\n2. Khi đăng nhập, ngoài Mật khẩu/RADIUS Auth, người dùng cần nhập mã OTP 6 chữ số. TotpService sẽ xác thực mã OTP trong khoảng thời gian hợp lệ (Window 30 giây).")

    # Section 3: Architecture & Domain Setup
    h3 = doc.add_heading("3. Đề Xuất Kiến Trúc Hạ Tầng & Triển Khai", level=1)
    h3.style.font.color.rgb = RGBColor(0, 51, 102)
    
    h3_1 = doc.add_heading("3.1. Cấu hình Tên miền & Điều hướng Routing", level=2)
    h3_1.style.font.color.rgb = RGBColor(11, 79, 108)
    
    p_dom = doc.add_paragraph()
    p_dom.add_run("Hệ thống đề xuất sử dụng mô hình ")
    r_dom_b = p_dom.add_run("Đồng nhất Tên miền (Single Domain Topology)")
    r_dom_b.bold = True
    p_dom.add_run(" để triệt tiêu hoàn toàn nguy cơ lỗi CORS và tối ưu hóa việc quản lý chứng chỉ SSL trong nội bộ ngân hàng:")
    
    dom_table = doc.add_table(rows=4, cols=3)
    dom_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(dom_table, "CCCCCC")
    headers = ["Thành phần", "Đường dẫn Tên miền (URL)", "Giao thức & Cổng"]
    for idx, text in enumerate(headers):
        cell = dom_table.rows[0].cells[idx]
        set_cell_background(cell, "003366")
        p_h = cell.paragraphs[0]
        r = p_h.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    routes = [
        ("Giao diện Frontend SPA", "https://hopdong.coopbank.vn/", "HTTPS - Port 443"),
        ("Backend RESTful API", "https://hopdong.coopbank.vn/api/...", "HTTPS Proxy -> Backend 8080"),
        ("SignalR Realtime Hub", "https://hopdong.coopbank.vn/hub/notifications", "WSS (WebSocket Secure)")
    ]
    for row_idx, data in enumerate(routes, start=1):
        row = dom_table.rows[row_idx]
        bg = "F9FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p_c = cell.paragraphs[0]
            p_c.paragraph_format.space_after = Pt(0)
            p_c.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    h3_2 = doc.add_heading("3.2. Phương án Hạ tầng Triển khai (Docker vs IIS)", level=2)
    h3_2.style.font.color.rgb = RGBColor(11, 79, 108)
    
    p_opt = doc.add_paragraph()
    p_opt.add_run("Đội dự án đề xuất 2 phương án tùy thuộc vào định hướng hạ tầng của Trung tâm CNTT:")
    
    p_o1 = doc.add_paragraph(style='List Bullet')
    r_o1 = p_o1.add_run("Phương án 1 (Khuyên dùng): Docker Container & Nginx Reverse Proxy\n")
    r_o1.bold = True
    p_o1.add_run("Backend chạy trong Docker Container (Image `qlda-backend:latest`). Nginx Reverse Proxy chịu trách nhiệm SSL Termination, tĩnh hóa file Frontend React và proxy kết nối WebSocket cho SignalR.")
    
    p_o2 = doc.add_paragraph(style='List Bullet')
    r_o2 = p_o2.add_run("Phương án 2: Windows Server & IIS (Internet Information Services)\n")
    r_o2.bold = True
    p_o2.add_run("Host Backend qua ASP.NET Core Module v2 (ANCM) trên IIS. Bật tính năng IIS WebSocket Protocol và sử dụng URL Rewrite Module cho Frontend SPA Routing.")

    # Section 4: Network Matrix
    h4 = doc.add_heading("4. Yêu Cầu Tiền Đề & Ma Trận Kết Nối Mạng (Network Matrix)", level=1)
    h4.style.font.color.rgb = RGBColor(0, 51, 102)
    
    p_net = doc.add_paragraph()
    p_net.add_run("Quản trị hệ thống và An toàn thông tin cần mở các luồng kết nối Firewall theo bảng thông số dưới đây trước thời điểm deploy:")
    
    net_table = doc.add_table(rows=5, cols=4)
    net_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(net_table, "CCCCCC")
    
    net_headers = ["Nguồn (Source)", "Đích (Destination)", "Cổng (Port/Protocol)", "Mục đích sử dụng"]
    for idx, text in enumerate(net_headers):
        cell = net_table.rows[0].cells[idx]
        set_cell_background(cell, "003366")
        p_h = cell.paragraphs[0]
        r = p_h.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    net_rules = [
        ("User Browser (Client)", "Reverse Proxy / Nginx", "TCP 443 (HTTPS)", "Truy cập ứng dụng Web"),
        ("App Server (Backend)", "MySQL DB Server (10.225.11.201)", "TCP 3306 (MySQL)", "Kết nối Cơ sở dữ liệu QLDA"),
        ("App Server (Backend)", "RADIUS Server (10.224.0.94)", "UDP 1812 (RADIUS)", "Xác thực tài khoản Domain nội bộ"),
        ("App Server (Backend)", "SMTP Mail Server", "TCP 25 / 587 (SMTP)", "Gửi mail cảnh báo hợp đồng hết hạn")
    ]
    for row_idx, data in enumerate(net_rules, start=1):
        row = net_table.rows[row_idx]
        bg = "F9FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p_c = cell.paragraphs[0]
            p_c.paragraph_format.space_after = Pt(0)
            p_c.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 5: Checklist
    h5 = doc.add_heading("5. Các Danh Mục Cần Đặc Biệt Lưu Ý (Crucial Checklist)", level=1)
    h5.style.font.color.rgb = RGBColor(0, 51, 102)
    
    add_callout(
        doc,
        title="1. CẤU HÌNH TẮT AUTOMIGRATE VÀ SEED DATA TRÊN PRODUCTION",
        text="Trên môi trường Production, tuyệt đối KHÔNG bật 'Database:AutoMigrate=true' hoặc 'Database:SeedSampleData=true' trong appsettings.json. Mọi thay đổi cấu trúc DB bắt buộc phải chạy qua Script SQL Idempotent do Đội DBA kiểm duyệt.",
        border_color="CC0000",
        bg_color="FFF0F0"
    )
    
    add_callout(
        doc,
        title="2. THAY ĐỔI CHUỖI KHÓA BẢO MẬT JWT (SECRET KEY)",
        text="Khóa SecretKey trong 'JwtSettings' trên Production phải là chuỗi ngẫu nhiên dài tối thiểu 256-bit (32+ ký tự), không được dùng lại chuỗi mặc định ở môi trường Development để tránh rủi ro giả mạo Token xác thực.",
        border_color="E65100",
        bg_color="FFF8E1"
    )

    add_callout(
        doc,
        title="3. CẤU HÌNH SIGNALR WEBSOCKET HANDSHAKE QUA REVERSE PROXY",
        text="Do kết nối WebSocket handshake của SignalR gửi token qua query string ('/hub/notifications?access_token=...'), Reverse Proxy (Nginx/IIS) phải được cấu hình 'proxy_set_header Upgrade $http_upgrade' và 'Connection \"Upgrade\"' để tránh lỗi gãy kết nối 401/404.",
        border_color="003366",
        bg_color="F0F4F8"
    )

    add_callout(
        doc,
        title="4. TẮT GIAO DIỆN SWAGGER UI VÀ BẬT LOG AUDIT REAL IP",
        text="Trên môi trường Production, ứng dụng đã cài đặt ẩn giao diện Swagger UI công khai. Đồng thời bắt buộc bật middleware 'UseForwardedHeaders' để thu thập chính xác IP thực của Client qua X-Forwarded-For cho mục đích Audit Log.",
        border_color="003366",
        bg_color="F0F4F8"
    )

    # Section 6: Steps
    h6 = doc.add_heading("6. Quy Trình Triển Khai Chi Tiết Từng Bước (Step-by-Step)", level=1)
    h6.style.font.color.rgb = RGBColor(0, 51, 102)
    
    steps = [
        ("Bước 1: Sao lưu & Cập nhật Cơ sở dữ liệu", "Đội DBA thực hiện Snapshot/Backup CSDL 'QLDA' trên server 10.225.11.201. Tiếp theo, thực thi tệp script SQL 'deploy_v1.0.0.sql' đã được kiểm duyệt để khởi tạo đầy đủ các bảng, chỉ mục và dữ liệu phân quyền ban đầu."),
        ("Bước 2: Cấu hình Tệp appsettings.Production.json", "Cập nhật chuỗi kết nối MySQL 'DefaultConnection', JWT SecretKey, thông số RADIUS Server (10.224.0.94) và địa chỉ SMTP Mail vào biến môi trường hoặc tệp appsettings.Production.json của ứng dụng."),
        ("Bước 3: Biên dịch & Khởi chạy Backend Service", "Thực hiện 'dotnet publish -c Release' hoặc build Docker Image 'coopbank/contract-backend:v1.0.0'. Khởi chạy service Backend và kiểm tra trạng thái liveness qua log hệ thống."),
        ("Bước 4: Deploy Frontend & Cấu hình Reverse Proxy", "Sao chép thư mục dist của Frontend React vào web root của Nginx/IIS. Cấu hình các rule routing /api, /hub và áp dụng chứng chỉ SSL Certificate nội bộ CoopBank."),
        ("Bước 5: Thực hiện Smoke Test & Nghiệm thu", "Đội QA/DEV tiến hành kiểm thử nhanh 10 kịch bản Smoke Test (Đăng nhập RADIUS, MFA TOTP, Tạo dự án/hợp đồng, Nhận thông báo SignalR, Gửi mail cảnh báo).")
    ]
    for title, detail in steps:
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_after = Pt(4)
        rs = p_s.add_run(title)
        rs.bold = True
        rs.font.color.rgb = RGBColor(11, 79, 108)
        
        p_sd = doc.add_paragraph()
        p_sd.paragraph_format.left_indent = Inches(0.2)
        p_sd.paragraph_format.space_after = Pt(8)
        p_sd.add_run(detail)

    # Section 7: Rollback
    h7 = doc.add_heading("7. Kế Hoạch Rollback (Phòng Ngừa Sự Cố)", level=1)
    h7.style.font.color.rgb = RGBColor(0, 51, 102)
    
    p_rb = doc.add_paragraph()
    p_rb.add_run("Trong trường hợp xảy ra sự cố nghiêm trọng không thể khắc phục trong thời gian cửa sổ triển khai (Deployment Window):")
    
    rb_items = [
        ("Bước 1:", "Dừng các service Backend và Reverse Proxy mới triển khai."),
        ("Bước 2:", "Khôi phục lại trạng thái CSDL 'QLDA' từ bản Snapshot/Backup ở Bước 1."),
        ("Bước 3:", "Kích hoạt lại phiên bản ứng dụng cũ (nếu có) hoặc hiển thị trang thông báo bảo trì tạm thời."),
        ("Bước 4:", "Họp đánh giá nguyên nhân sự cố (Post-mortem) và lập kế hoạch khắc phục.")
    ]
    for b, t in rb_items:
        p_rbi = doc.add_paragraph(style='List Bullet')
        p_rbi.paragraph_format.space_after = Pt(3)
        rb_b = p_rbi.add_run(b + " ")
        rb_b.bold = True
        p_rbi.add_run(t)

    # Section 8: Sign-off Table
    h8 = doc.add_heading("8. Bảng Ký Duyệt Triển Khai (Sign-off & Approvals)", level=1)
    h8.style.font.color.rgb = RGBColor(0, 51, 102)
    
    p_so = doc.add_paragraph()
    p_so.add_run("Bản tài liệu này được sự đồng thuận và phê duyệt của đại diện các bộ phận liên quan:")
    
    sign_table = doc.add_table(rows=3, cols=4)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(sign_table, "CCCCCC")
    
    sign_headers = ["Đội Phát Triển (DEV)", "Quản Trị Hạ Tầng (Ops)", "An Toàn Thông Tin (Sec)", "Lãnh Đạo TT CNTT"]
    for idx, text in enumerate(sign_headers):
        cell = sign_table.rows[0].cells[idx]
        set_cell_background(cell, "F0F4F8")
        p_h = cell.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_h.add_run(text)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0, 51, 102)
        
    for col_idx in range(4):
        cell_sig = sign_table.rows[1].cells[col_idx]
        set_cell_margins(cell_sig, top=400, bottom=400, left=50, right=50)
        p_sig = cell_sig.paragraphs[0]
        p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_s = p_sig.add_run("(Ký & Họ tên)")
        r_s.font.italic = True
        r_s.font.size = Pt(8.5)
        r_s.font.color.rgb = RGBColor(150, 150, 150)
        
        cell_date = sign_table.rows[2].cells[col_idx]
        set_cell_margins(cell_date, top=50, bottom=50, left=50, right=50)
        p_d = cell_date.paragraphs[0]
        p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_d = p_d.add_run("Ngày: ...../...../2026")
        r_d.font.size = Pt(8.5)
        
    output_path = r"c:\Users\mdqa7\Documents\CoopBank_Project\Contract_Management\Tai_Lieu_Yeu_Cau_Deploy_Va_Dac_Ta_Ky_Thuat_CoopBank.docx"
    doc.save(output_path)
    print(f"Successfully updated docx file with Technical Specs at: {output_path}")

if __name__ == "__main__":
    build_docx()
